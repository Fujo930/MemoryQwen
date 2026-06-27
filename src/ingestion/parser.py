"""
MemoryQwen — DocumentParser 文件解析器
支持 txt, md, pdf, docx, 代码文件
"""

from __future__ import annotations

import os
import logging
from pathlib import Path
from dataclasses import dataclass
from typing import Any

logger = logging.getLogger(__name__)


class UnsupportedFormatError(Exception):
    """不支持的格式"""
    pass


@dataclass
class Document:
    """解析后的文档"""
    file_path: str
    file_type: str
    title: str
    content: str
    metadata: dict
    pages: int | None = None
    char_count: int = 0


class DocumentParser:
    """多格式文件解析器"""

    def __init__(self, config: Any | None = None):
        self.config = config

    def can_handle(self, file_path: str) -> bool:
        """检查是否支持该格式"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions()

    def supported_extensions(self) -> list[str]:
        """返回支持的扩展名列表"""
        return [".txt", ".md", ".pdf", ".docx", ".py", ".js", ".ts",
                ".java", ".cpp", ".rs", ".go", ".c", ".h", ".html",
                ".css", ".json", ".yaml", ".yml", ".toml", ".xml",
                ".sh", ".bat", ".ps1", ".sql", ".r", ".rb", ".php"]

    async def parse(self, file_path: str) -> Document:
        """解析文件"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        title = path.stem

        if ext == ".txt":
            content = await self._parse_txt(path)
        elif ext == ".md":
            content = await self._parse_md(path)
        elif ext == ".pdf":
            content = await self._parse_pdf(path)
        elif ext == ".docx":
            content = await self._parse_docx(path)
        elif ext in self.supported_extensions():
            content = await self._parse_code(path)
        else:
            raise UnsupportedFormatError(f"不支持的格式: {ext}")

        return Document(
            file_path=str(path),
            file_type=ext.lstrip("."),
            title=title,
            content=content,
            metadata={
                "source_path": str(path),
                "source_type": ext.lstrip("."),
                "doc_title": title,
                "file_size": path.stat().st_size,
            },
            char_count=len(content),
        )

    async def _parse_txt(self, path: Path) -> str:
        """解析 txt 文件"""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    async def _parse_md(self, path: Path) -> str:
        """解析 markdown 文件（保留纯文本）"""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return content  # 保留 markdown 格式，chunker 会利用标题层级

    async def _parse_pdf(self, path: Path) -> str:
        """解析 PDF 文件"""
        try:
            import pypdf
            text_parts = []
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, trying fallback")
            return f"[PDF file: {path.name} — install pypdf to extract text]"
        except Exception as e:
            logger.error("PDF parse failed: %s", e)
            return f"[PDF parse error: {e}]"

    async def _parse_docx(self, path: Path) -> str:
        """解析 docx 文件"""
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed")
            return f"[DOCX file: {path.name} — install python-docx to extract text]"
        except Exception as e:
            logger.error("DOCX parse failed: %s", e)
            return f"[DOCX parse error: {e}]"

    async def _parse_code(self, path: Path) -> str:
        """解析代码文件"""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
